"""Validate structural requirements recovered from the competition template."""

from __future__ import annotations

import hashlib
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_TEMPLATE = ROOT / "最新版！2026数学建模国赛标准论文Word模板.doc"
PAPER = ROOT / "output" / "docx" / "competition_paper_draft.docx"
EXPECTED_TEMPLATE_HASH = "82DC4197C4297CC9F166475DD140720C7DE74BECC4B2D9C350C7D6ADBB389A77"
REQUIRED = [
    "摘要", "1 问题重述", "2 问题分析", "3 模型假设", "4 符号说明",
    "5 模型建立与求解", "6 模型检验", "7 模型评价与改进",
    "8 AI 工具使用声明", "参考文献", "附录",
]


def main() -> int:
    failures: list[str] = []
    actual_hash = hashlib.sha256(SOURCE_TEMPLATE.read_bytes()).hexdigest().upper()
    if actual_hash != EXPECTED_TEMPLATE_HASH:
        failures.append("the uploaded binary template changed after recovery")
    if not PAPER.exists():
        failures.append("competition paper DOCX is missing")
    else:
        doc = Document(PAPER)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        failures.extend(f"missing required section: {item}" for item in REQUIRED if item not in text)
        if "目录" in text:
            failures.append("template forbids a table of contents")
        if any(token in text for token in ("XXX", "占位", "待补")):
            failures.append("placeholder text remains in the paper")
        if len(doc.tables) < 5:
            failures.append(f"expected at least 5 result tables, found {len(doc.tables)}")
        if len(doc.inline_shapes) < 10:
            failures.append(f"expected at least 10 figures, found {len(doc.inline_shapes)}")
        with zipfile.ZipFile(PAPER) as package:
            document_xml = package.read("word/document.xml")
            if document_xml.count(b"<m:oMath") < 8:
                failures.append("editable equation count is unexpectedly low")
    if failures:
        for failure in failures:
            print("FAIL:", failure)
        return 1
    print("Competition paper validation passed: structure, figures, tables, equations, and template hash.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
