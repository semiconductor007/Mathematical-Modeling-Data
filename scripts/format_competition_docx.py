"""Apply the recovered competition-template rules to a Pandoc DOCX draft."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, east_asia: str, western: str, size: float, bold: bool | None = None) -> None:
    run.font.name = western
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_border(parent, edge: str, val: str, size: int = 8) -> None:
    tag = "w:tblBorders" if parent.tag == qn("w:tblPr") else "w:tcBorders"
    borders = parent.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        parent.append(borders)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:color"), "000000")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    set_run_font(run, "宋体", "Times New Roman", 9)


def all_paragraphs(doc: Document):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def style_tables(doc: Document, usable_dxa: int = 9354) -> None:
    for table in doc.tables:
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        for edge in ("top", "bottom"):
            set_border(tbl_pr, edge, "single", 12)
        for edge in ("left", "right", "insideH", "insideV"):
            set_border(tbl_pr, edge, "nil", 0)

        ncols = len(table.columns)
        raw = []
        for col in range(ncols):
            length = max((len(row.cells[col].text.strip()) for row in table.rows), default=1)
            raw.append(min(5.0, max(1.0, length ** 0.55)))
        widths = [int(usable_dxa * value / sum(raw)) for value in raw]
        widths[-1] += usable_dxa - sum(widths)
        for col_index, width in enumerate(widths):
            table.columns[col_index].width = Cm(width / 567.0)

        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                tr_pr = row._tr.get_or_add_trPr()
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "true")
                tr_pr.append(repeat)
            for col_index, cell in enumerate(row.cells):
                cell.width = Cm(widths[col_index] / 567.0)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                if row_index == 0:
                    set_border(tc_pr, "bottom", "single", 8)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.left_indent = Pt(0)
                    paragraph.paragraph_format.right_indent = Pt(0)
                    paragraph.alignment = (WD_ALIGN_PARAGRAPH.LEFT if col_index == 1
                                           else WD_ALIGN_PARAGRAPH.CENTER)
                    for run in paragraph.runs:
                        set_run_font(run, "宋体", "Times New Roman", 10.5, row_index == 0)


def main(source: Path, output: Path) -> None:
    doc = Document(source)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)
        section.different_first_page_header_footer = False
        for footer in (section.footer, section.even_page_footer, section.first_page_footer):
            for p in footer.paragraphs:
                p.clear()
            add_page_number(footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    style_specs = {
        "Title": ("黑体", 18, True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12),
        "Heading 1": ("黑体", 14, True, WD_ALIGN_PARAGRAPH.LEFT, 12, 6),
        "Heading 2": ("黑体", 12, True, WD_ALIGN_PARAGRAPH.LEFT, 10, 4),
        "Heading 3": ("黑体", 12, True, WD_ALIGN_PARAGRAPH.LEFT, 8, 3),
        "Caption": ("宋体", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, 3, 8),
    }
    for name, (font, size, bold, align, before, after) in style_specs.items():
        if name not in styles:
            continue
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.paragraph_format.alignment = align
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    in_references = False
    for paragraph in all_paragraphs(doc):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if text == "参考文献":
            in_references = True
        elif text == "附录":
            in_references = False
        if style_name == "Title":
            paragraph.paragraph_format.keep_with_next = True
        if text in {"摘要", "参考文献", "附录"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text.startswith("关键词："):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
        if re.match(r"^图\s*\d+", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(8)
            for run in paragraph.runs:
                set_run_font(run, "宋体", "Times New Roman", 10.5, False)
                run.italic = False
        if in_references and re.match(r"^\[\d+\]", text):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if style_name not in {"Title", "Heading 1", "Heading 2", "Heading 3", "Caption"}:
            paragraph.paragraph_format.widow_control = True
        for run in paragraph.runs:
            if style_name == "Title":
                set_run_font(run, "黑体", "Times New Roman", 18, True)
            elif style_name.startswith("Heading"):
                size = 14 if style_name == "Heading 1" else 12
                set_run_font(run, "黑体", "Times New Roman", size, True)
            elif re.match(r"^图\s*\d+", text):
                set_run_font(run, "宋体", "Times New Roman", 10.5, False)
                run.italic = False
            elif style_name == "Caption":
                set_run_font(run, "宋体", "Times New Roman", 10.5, False)
            elif not paragraph._parent.__class__.__name__.endswith("Cell"):
                set_run_font(run, "宋体", "Times New Roman", 12)

    style_tables(doc)
    max_width = Cm(15.6)
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    # Word computes page fields during normal pagination; avoid an open-time
    # refresh of every field because large embedded figures can stall automation.
    update.set(qn("w:val"), "false")
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Formatted {output}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: format_competition_docx.py INPUT.docx OUTPUT.docx")
    main(Path(sys.argv[1]), Path(sys.argv[2]))
